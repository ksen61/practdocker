from rest_framework import viewsets, generics, status, filters
from rest_framework.decorators import api_view, action, authentication_classes, permission_classes
from rest_framework.response import Response
from rest_framework.permissions import AllowAny, IsAuthenticated
from rest_framework.pagination import PageNumberPagination
from django.contrib.auth import authenticate, login as django_login, logout as django_logout, update_session_auth_hash
from django.core.validators import validate_email
from django.core.exceptions import ValidationError
from django.contrib.auth.models import User
from django.db import models
from django.db.models import Min, Max, OuterRef, Subquery, Count, Exists
from django.db.models.functions import Coalesce
from django.utils import timezone
from django.utils.html import strip_tags
from django.utils.crypto import get_random_string
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.password_validation import validate_password
import json
from django.http import HttpResponse
import io
from django.conf import settings
from django.core.mail import EmailMultiAlternatives
from django.utils.dateparse import parse_date

from documentflow.models import (
    DocumentType, Department, Document, DocumentStatus, DocumentRouteTemplate,
    User, Approval, DocumentFile, DocumentVersion, Notification, EmailChangeRequest,
    Replacement, EmployeeStatus
)
from .serializers import (
    DocumentTypeSerializer, DepartmentSerializer,
    UserSerializer, DocumentSerializer,
    DocumentStatusSerializer, DocumentCreateSerializer,
    DocumentRouteSerializer, DocumentDetailSerializer
)


# ============ PAGINATION ============


def _is_current_step(document, approval):
    if document.approval_order != 'sequential':
        return True
    min_step = Approval.objects.filter(
        document=document,
        decision='pending',
        cycle=approval.cycle
    ).aggregate(min_step=Min('step'))['min_step']
    return min_step is None or approval.step == min_step


class StandardPagination(PageNumberPagination):
    page_size = 100
    page_size_query_param = 'page_size'
    max_page_size = 1000


def _resolve_approver(user):
    if user.status in [
        EmployeeStatus.VACATION,
        EmployeeStatus.SICK,
        EmployeeStatus.BUSINESS_TRIP,
        EmployeeStatus.MATERNITY,
        EmployeeStatus.IDLE,
        EmployeeStatus.OTHER,
        EmployeeStatus.DISMISSED,
    ]:
        today = timezone.localdate()
        replacement = (
            Replacement.objects
            .filter(
                absent_employee=user,
                is_active=True,
                start_date__lte=today,
                end_date__gte=today,
                replacement_employee__is_active=True,
                replacement_employee__status=EmployeeStatus.WORKING
            )
            .order_by('-start_date')
            .first()
        )
        if replacement:
            return replacement.replacement_employee
        if user.department and user.department.head:
            head = user.department.head
            if head.is_active and head.status == EmployeeStatus.WORKING and head.id != user.id:
                return head
        fallback = (
            User.objects
            .filter(
                department=user.department,
                is_active=True,
                status=EmployeeStatus.WORKING
            )
            .exclude(id=user.id)
            .order_by('last_name', 'first_name')
            .first()
        )
        if fallback:
            return fallback
        return None
    return user


# ============ VIEWSETS ============


class DepartmentViewSet(viewsets.ModelViewSet):
    """Отделы"""
    queryset = Department.objects.all()
    serializer_class = DepartmentSerializer
    permission_classes = [AllowAny]
    pagination_class = None


# ============ AUTH ENDPOINTS ============


@csrf_exempt
@api_view(['POST'])
@authentication_classes([])
@permission_classes([AllowAny])
def LoginAPIView(request):
    """
    POST /api/login/ - Вход
    
    Требует:
    {
        "username": "user@example.com",
        "password": "password123"
    }
    
    Возвращает:
    {
        "token": "abc123def456",
        "user": {
            "id": 1,
            "username": "user@example.com",
            "email": "user@example.com",
            "full_name": "John Doe"
        }
    }
    """
    username = request.data.get('username')
    password = request.data.get('password')
    
    if not username or not password:
        return Response({
            'error': 'Требуются username и password'
        }, status=status.HTTP_400_BAD_REQUEST)
    
    user = authenticate(username=username, password=password)
    
    if not user:
        return Response({
            'error': 'Неверные учетные данные'
        }, status=status.HTTP_401_UNAUTHORIZED)
    
    django_login(request, user)

    return Response({
        'user': {
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'full_name': f"{user.first_name} {user.last_name}".strip() or user.username
        }
        ,
        'must_change_password': getattr(user, 'must_change_password', False)
    }, status=status.HTTP_200_OK)


@csrf_exempt
@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def LogoutAPIView(request):
    """POST /api/logout/ - Выход"""
    if request.user.is_authenticated:
        django_logout(request)
    
    return Response({
        'status': 'success',
        'message': 'Вы вышли из системы'
    }, status=status.HTTP_200_OK)


@api_view(['GET'])
@authentication_classes([])
@permission_classes([AllowAny])
def CheckAuthAPIView(request):
    """GET /api/check-auth/ - Проверка авторизации"""
    return Response({
        'authenticated': request.user.is_authenticated,
        'user_id': request.user.id if request.user.is_authenticated else None,
        'must_change_password': getattr(request.user, 'must_change_password', False) if request.user.is_authenticated else False
    })


@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def ChangePasswordAPIView(request):
    """
    POST /api/password-change/ - Смена пароля
    """
    current_password = request.data.get('current_password')
    new_password = request.data.get('new_password')
    confirm_password = request.data.get('confirm_password')

    if not current_password or not new_password or not confirm_password:
        return Response({'error': 'Все поля обязательны'}, status=status.HTTP_400_BAD_REQUEST)

    if not request.user.check_password(current_password):
        return Response({'error': 'Текущий пароль неверный'}, status=status.HTTP_400_BAD_REQUEST)

    if new_password != confirm_password:
        return Response({'error': 'Пароли не совпадают'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        validate_password(new_password, user=request.user)
    except ValidationError as exc:
        return Response({'error': '; '.join(exc.messages)}, status=status.HTTP_400_BAD_REQUEST)

    request.user.set_password(new_password)
    if hasattr(request.user, 'must_change_password'):
        request.user.must_change_password = False
    request.user.save(update_fields=['password', 'must_change_password'] if hasattr(request.user, 'must_change_password') else ['password'])
    update_session_auth_hash(request, request.user)

    return Response({'status': 'success'}, status=status.HTTP_200_OK)


@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def RequestEmailChangeAPIView(request):
    """
    POST /api/email-change/request/ - Запрос смены email (код приходит на старую почту)
    """
    password = request.data.get('password')
    new_email = request.data.get('new_email')

    if not password or not new_email:
        return Response({'error': 'Все поля обязательны'}, status=status.HTTP_400_BAD_REQUEST)

    if not request.user.check_password(password):
        return Response({'error': 'Пароль неверный'}, status=status.HTTP_400_BAD_REQUEST)

    if not request.user.email:
        return Response({'error': 'У пользователя не указан текущий email'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        validate_email(new_email)
    except ValidationError:
        return Response({'error': 'Некорректный email'}, status=status.HTTP_400_BAD_REQUEST)

    if new_email.strip().lower() == request.user.email.strip().lower():
        return Response({'error': 'Новый email совпадает с текущим'}, status=status.HTTP_400_BAD_REQUEST)

    EmailChangeRequest.objects.filter(user=request.user).delete()
    code = get_random_string(6, allowed_chars='0123456789')
    EmailChangeRequest.objects.create(user=request.user, new_email=new_email, code=code)

    subject = 'Подтверждение смены email'
    html_message = f"""
    <div style="font-family: Arial, sans-serif; background:#f6f7fb; padding:24px;">
        <div style="max-width:640px; margin:0 auto; background:#ffffff; border-radius:14px; border:1px solid #e6e8ef; overflow:hidden;">
            <div style="padding:18px 22px; background:#0f766e;">
                <div style="color:#ffffff; font-size:18px; font-weight:700;">DocumentFlow</div>
                <div style="color:#d1fae5; font-size:12px;">Подтверждение смены email</div>
            </div>
            <div style="padding:22px;">
                <p style="margin:0 0 10px 0; color:#111827; font-size:14px;">Здравствуйте!</p>
                <p style="margin:0 0 16px 0; color:#374151; font-size:13px;">
                    Вы запросили смену email в системе DocumentFlow. Используйте код ниже, чтобы подтвердить действие.
                </p>
                <div style="border:1px solid #e5e7eb; border-radius:10px; padding:14px; background:#f9fafb; text-align:center;">
                    <div style="color:#6b7280; font-size:11px; text-transform:uppercase; letter-spacing:0.08em;">Код подтверждения</div>
                    <div style="font-size:22px; font-weight:700; color:#111827; letter-spacing:0.2em;">{code}</div>
                </div>
                <p style="margin:16px 0 0 0; color:#6b7280; font-size:12px;">
                    Код действует 10 минут. Если вы не запрашивали смену email — просто проигнорируйте это письмо.
                </p>
            </div>
            <div style="padding:12px 22px; background:#f3f4f6; color:#9ca3af; font-size:11px;">
                Это письмо отправлено автоматически, отвечать на него не нужно.
            </div>
        </div>
    </div>
    """
    message = strip_tags(html_message)
    try:
        email = EmailMultiAlternatives(
            subject=subject,
            body=message,
            from_email=getattr(settings, 'DEFAULT_FROM_EMAIL', None),
            to=[request.user.email],
        )
        email.attach_alternative(html_message, "text/html")
        email.send(fail_silently=False)
    except Exception as exc:
        return Response({'error': f'Не удалось отправить письмо: {exc}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    return Response({'status': 'code_sent'}, status=status.HTTP_200_OK)


@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def ConfirmEmailChangeAPIView(request):
    """
    POST /api/email-change/confirm/ - Подтверждение смены email
    """
    code = request.data.get('code')
    if not code:
        return Response({'error': 'Введите код'}, status=status.HTTP_400_BAD_REQUEST)

    req = EmailChangeRequest.objects.filter(user=request.user).order_by('-created_at').first()
    if not req:
        return Response({'error': 'Запрос на смену email не найден'}, status=status.HTTP_400_BAD_REQUEST)

    if req.is_expired():
        req.delete()
        return Response({'error': 'Код просрочен. Запросите новый.'}, status=status.HTTP_400_BAD_REQUEST)

    if req.code != str(code).strip():
        req.attempts += 1
        req.save(update_fields=['attempts'])
        if req.attempts >= 5:
            req.delete()
            return Response({'error': 'Превышено число попыток. Запросите новый код.'}, status=status.HTTP_400_BAD_REQUEST)
        return Response({'error': 'Неверный код'}, status=status.HTTP_400_BAD_REQUEST)

    request.user.email = req.new_email
    request.user.save(update_fields=['email'])
    req.delete()

    return Response({'status': 'success', 'email': request.user.email}, status=status.HTTP_200_OK)


@csrf_exempt
@api_view(['POST'])
@permission_classes([IsAuthenticated])
def ResendEmailChangeCodeAPIView(request):
    """
    POST /api/email-change/resend/ - Повторная отправка кода
    """
    req = EmailChangeRequest.objects.filter(user=request.user).order_by('-created_at').first()
    if not req:
        return Response({'error': 'Запрос на смену email не найден'}, status=status.HTTP_400_BAD_REQUEST)

    if req.is_expired():
        req.delete()
        return Response({'error': 'Код просрочен. Запросите новый.'}, status=status.HTTP_400_BAD_REQUEST)

    if req.resend_count >= 3:
        return Response({'error': 'Достигнут лимит повторных отправок.'}, status=status.HTTP_400_BAD_REQUEST)

    now = timezone.now()
    if req.last_sent_at and (now - req.last_sent_at).total_seconds() < 60:
        return Response({'error': 'Повторная отправка доступна через 60 секунд.'}, status=status.HTTP_400_BAD_REQUEST)

    req.resend_count += 1
    req.last_sent_at = now
    req.save(update_fields=['resend_count', 'last_sent_at'])

    subject = 'Подтверждение смены email'
    html_message = f"""
    <div style="font-family: Arial, sans-serif; background:#f6f7fb; padding:24px;">
        <div style="max-width:640px; margin:0 auto; background:#ffffff; border-radius:14px; border:1px solid #e6e8ef; overflow:hidden;">
            <div style="padding:18px 22px; background:#0f766e;">
                <div style="color:#ffffff; font-size:18px; font-weight:700;">DocumentFlow</div>
                <div style="color:#d1fae5; font-size:12px;">Повторная отправка кода</div>
            </div>
            <div style="padding:22px;">
                <p style="margin:0 0 10px 0; color:#111827; font-size:14px;">Здравствуйте!</p>
                <p style="margin:0 0 16px 0; color:#374151; font-size:13px;">
                    Вы запросили повторную отправку кода для смены email.
                </p>
                <div style="border:1px solid #e5e7eb; border-radius:10px; padding:14px; background:#f9fafb; text-align:center;">
                    <div style="color:#6b7280; font-size:11px; text-transform:uppercase; letter-spacing:0.08em;">Код подтверждения</div>
                    <div style="font-size:22px; font-weight:700; color:#111827; letter-spacing:0.2em;">{req.code}</div>
                </div>
                <p style="margin:16px 0 0 0; color:#6b7280; font-size:12px;">
                    Код действует 10 минут. Если вы не запрашивали смену email — просто проигнорируйте это письмо.
                </p>
            </div>
            <div style="padding:12px 22px; background:#f3f4f6; color:#9ca3af; font-size:11px;">
                Это письмо отправлено автоматически, отвечать на него не нужно.
            </div>
        </div>
    </div>
    """
    message = strip_tags(html_message)
    try:
        email = EmailMultiAlternatives(
            subject=subject,
            body=message,
            from_email=getattr(settings, 'DEFAULT_FROM_EMAIL', None),
            to=[request.user.email],
        )
        email.attach_alternative(html_message, "text/html")
        email.send(fail_silently=False)
    except Exception as exc:
        return Response({'error': f'Не удалось отправить письмо: {exc}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    return Response({'status': 'resent'}, status=status.HTTP_200_OK)


@api_view(['GET', 'POST'])
@permission_classes([IsAuthenticated])
def ReplacementSelfAPIView(request):
    if request.method == 'GET':
        replacements = Replacement.objects.filter(absent_employee=request.user).order_by('-start_date')
        data = []
        for r in replacements:
            data.append({
                'id': r.id,
                'replacement_employee': {
                    'id': r.replacement_employee.id,
                    'display_name': r.replacement_employee.full_name,
                },
                'reason': r.reason,
                'reason_display': r.get_reason_display(),
                'start_date': r.start_date.isoformat() if r.start_date else None,
                'end_date': r.end_date.isoformat() if r.end_date else None,
                'is_active': r.is_active,
            })
        return Response({'results': data})

    replacement_id = request.data.get('replacement_employee')
    reason = request.data.get('reason')
    start_date = parse_date(request.data.get('start_date', ''))
    end_date = parse_date(request.data.get('end_date', ''))

    if reason == 'dismissed' and not start_date:
        start_date = timezone.localdate()

    if not replacement_id or not reason or not start_date or (not end_date and reason != 'dismissed'):
        return Response({'error': 'Все поля обязательны (дата окончания не требуется только при "Уволен")'}, status=status.HTTP_400_BAD_REQUEST)

    if end_date and start_date > end_date:
        return Response({'error': 'Дата окончания не может быть раньше даты начала'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        replacement_employee = User.objects.get(id=replacement_id, is_active=True)
    except User.DoesNotExist:
        return Response({'error': 'Сотрудник не найден'}, status=status.HTTP_400_BAD_REQUEST)

    if replacement_employee.id == request.user.id:
        return Response({'error': 'Нельзя назначить замену самому себе'}, status=status.HTTP_400_BAD_REQUEST)

    if replacement_employee.status in [EmployeeStatus.VACATION, EmployeeStatus.SICK, EmployeeStatus.DISMISSED]:
        return Response({'error': 'Сотрудник недоступен для замены'}, status=status.HTTP_400_BAD_REQUEST)

    replacement = Replacement.objects.create(
        absent_employee=request.user,
        replacement_employee=replacement_employee,
        reason=reason,
        start_date=start_date,
        end_date=end_date,
        created_by=request.user
    )

    return Response({
        'id': replacement.id,
        'status': 'success'
    }, status=status.HTTP_201_CREATED)


@api_view(['GET'])
@authentication_classes([])
@permission_classes([AllowAny])
def MeAPIView(request):
    """GET /api/me/ - Текущий пользователь"""
    if request.user.is_authenticated:
        return Response({
            'id': request.user.id,
            'username': request.user.username,
            'email': request.user.email,
            'full_name': f"{request.user.first_name} {request.user.last_name}".strip() or request.user.username
        })
    return Response({
        'error': 'Not authenticated'
    }, status=status.HTTP_401_UNAUTHORIZED)


# ============ LIST API VIEWS ============


class DocumentTypeListAPIView(generics.ListAPIView):
    """GET /api/document-types/ - Все типы документов"""
    queryset = DocumentType.objects.all()
    serializer_class = DocumentTypeSerializer
    permission_classes = [AllowAny]
    pagination_class = None


class UserListAPIView(generics.ListAPIView):
    """GET /api/users/ - Все активные пользователи"""
    queryset = User.objects.filter(is_active=True).exclude(status=EmployeeStatus.DISMISSED)
    serializer_class = UserSerializer
    permission_classes = [AllowAny]
    pagination_class = None


class DepartmentListAPIView(generics.ListAPIView):
    """GET /api/departments/ - Все отделы"""
    queryset = Department.objects.all()
    serializer_class = DepartmentSerializer
    permission_classes = [AllowAny]
    pagination_class = None


class DocumentStatusListAPIView(generics.ListAPIView):
    """GET /api/document-statuses/ - Все статусы документов"""
    queryset = DocumentStatus.objects.all()
    serializer_class = DocumentStatusSerializer
    permission_classes = [AllowAny]
    pagination_class = None


class DocumentRouteListAPIView(generics.ListAPIView):
    """
    GET /api/routes/ - Все маршруты
    GET /api/routes/?document_type=1 - Маршруты для типа документа
    """
    serializer_class = DocumentRouteSerializer
    permission_classes = [AllowAny]
    pagination_class = None

    def get_queryset(self):
        """Фильтруем по document_type если передан параметр"""
        queryset = DocumentRouteTemplate.objects.filter(is_active=True)

        document_type_id = self.request.query_params.get('document_type')
        if document_type_id:
            queryset = queryset.filter(document_type_id=document_type_id)

        return queryset


# ============ DOCUMENT CRUD ============


class DocumentListCreateAPIView(generics.ListCreateAPIView):
    queryset = Document.objects.all()
    permission_classes = [AllowAny]

    def get_serializer_class(self):
        return DocumentCreateSerializer if self.request.method == 'POST' else DocumentSerializer

    def perform_create(self, serializer):
        document = serializer.save(author=self.request.user)




class DocumentRetrieveUpdateDestroyAPIView(generics.RetrieveUpdateDestroyAPIView):
    """
    GET /api/documents/<id>/ - Получить документ
    PATCH /api/documents/<id>/ - Обновить документ
    DELETE /api/documents/<id>/ - Удалить документ
    """
    serializer_class = DocumentSerializer
    permission_classes = [AllowAny]

    def get_serializer_class(self):
        if self.request.method == 'GET':
            return DocumentDetailSerializer
        return DocumentSerializer

    def get_queryset(self):
        """Пользователи видят только свои документы"""
        user = self.request.user
        if user.is_authenticated:
            return Document.objects.filter(
                models.Q(author=user)
                | models.Q(responsible=user)
                | models.Q(approvals__approver=user)
            ).distinct()
        return Document.objects.none()


class DocumentFileCreateAPIView(generics.CreateAPIView):
    """POST /api/documents/<id>/files/ - Загрузить файл к документу"""
    serializer_class = DocumentSerializer
    permission_classes = [AllowAny]


# ============ INCOMING DOCUMENTS VIEW (Полученные документы) ============


class IncomingDocumentsAPIView(generics.ListAPIView):
    """
    GET /api/documents/incoming/ - Полученные документы (требующие согласования)
    Фильтры:
    - status: статус документа
    - priority: приоритет
    - document_type: тип документа
    - search: поиск по названию или номеру
    """
    serializer_class = DocumentSerializer
    permission_classes = [AllowAny]
    pagination_class = StandardPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['title', 'registration_number', 'description']
    ordering_fields = ['created_at', 'deadline', 'priority']
    ordering = ['-created_at']

    def get_queryset(self):
        """Получаем документы, на которых текущий пользователь является согласующим (Approval)"""
        user = self.request.user
        if not user.is_authenticated:
            return Document.objects.none()

        _ensure_cycle()

        doc_max_cycle = (
            Approval.objects.filter(document_id=OuterRef('pk'))
            .values('document_id')
            .annotate(max_cycle=Max('cycle'))
            .values('max_cycle')
        )
        approval_max_cycle = (
            Approval.objects.filter(document_id=OuterRef('document_id'))
            .values('document_id')
            .annotate(max_cycle=Max('cycle'))
            .values('max_cycle')
        )
        approval_min_pending_step = (
            Approval.objects.filter(
                document_id=OuterRef('document_id'),
                decision='pending',
                cycle=Subquery(approval_max_cycle),
            )
            .values('document_id')
            .annotate(min_step=Min('step'))
            .values('min_step')
        )
        current_pending_ids = (
            Approval.objects.filter(
                approver=user,
                decision='pending',
                cycle=Subquery(approval_max_cycle),
            )
            .filter(
                models.Q(step=Subquery(approval_min_pending_step))
                | models.Q(document__approval_order='parallel')
            )
            .values('document_id')
        )

        documents = (
            Document.objects
            .filter(is_archived=False)
            .filter(
                models.Q(id__in=Subquery(current_pending_ids))
                | models.Q(
                    approvals__approver=user,
                    approvals__decision__in=[
                        'approved', 'rejected', 'acknowledged', 'executed', 'returned'
                    ],
                    approvals__cycle=Subquery(doc_max_cycle),
                )
            )
            .select_related('document_type', 'status', 'author', 'responsible')
            .distinct()
        )

        # Применяем фильтры
        status_filter = self.request.query_params.get('status')
        if status_filter:
            documents = documents.filter(status_id=status_filter)

        priority_filter = self.request.query_params.get('priority')
        if priority_filter:
            documents = documents.filter(priority=priority_filter)

        document_type_filter = self.request.query_params.get('document_type')
        if document_type_filter:
            documents = documents.filter(document_type_id=document_type_filter)

        return documents


class MyDocumentsAPIView(generics.ListAPIView):
    """
    GET /api/documents/my/ - Мои документы (созданные мной)
    Фильтры:
    - status: статус документа
    - priority: приоритет
    - document_type: тип документа
    """
    serializer_class = DocumentSerializer
    permission_classes = [AllowAny]
    pagination_class = StandardPagination
    filter_backends = [filters.SearchFilter, filters.OrderingFilter]
    search_fields = ['title', 'registration_number', 'description']
    ordering_fields = ['created_at', 'deadline', 'priority']
    ordering = ['-created_at']

    def get_queryset(self):
        """Получаем документы, созданные текущим пользователем"""
        user = self.request.user
        if not user.is_authenticated:
            return Document.objects.none()

        documents = Document.objects.filter(
            author=user
        ).select_related('document_type', 'status', 'author', 'responsible')

        archived_filter = self.request.query_params.get('archived')
        if archived_filter is not None:
            if archived_filter in ['1', 'true', 'True', 'yes']:
                documents = documents.filter(is_archived=True)
            else:
                documents = documents.filter(is_archived=False)
        else:
            documents = documents.filter(is_archived=False)

        # Применяем фильтры
        status_filter = self.request.query_params.get('status')
        if status_filter:
            documents = documents.filter(status_id=status_filter)

        priority_filter = self.request.query_params.get('priority')
        if priority_filter:
            documents = documents.filter(priority=priority_filter)

        document_type_filter = self.request.query_params.get('document_type')
        if document_type_filter:
            documents = documents.filter(document_type_id=document_type_filter)

        return documents


# ============ DOCUMENT APPROVAL VIEWSET ============


class DocumentApprovalViewSet(viewsets.ModelViewSet):
    """
    ViewSet для операций с документами и их согласованием
    
    POST /api/documents/<id>/approve/ - Согласовать документ
    POST /api/documents/<id>/reject/ - Отклонить документ
    """
    queryset = Document.objects.all()
    serializer_class = DocumentSerializer
    permission_classes = [AllowAny]

    @action(detail=True, methods=['post'])
    def approve(self, request, pk=None):
        """Согласовать документ"""
        document = self.get_object()
        user = request.user

        try:
            approval = Approval.objects.get(
                document=document,
                approver=user,
                decision='pending'
            )
            if not _is_current_step(document, approval):
                return Response({
                    'status': 'error',
                    'message': 'Сейчас не ваш шаг согласования'
                }, status=status.HTTP_400_BAD_REQUEST)
            approval.decision = 'approved'
            approval.decided_at = timezone.now()
            approval.comment = request.data.get('comment', '')
            approval.save()
            _maybe_set_actual_deadline(document)

            return Response({
                'status': 'success',
                'message': f'Документ {document.registration_number} согласован'
            })
        except Approval.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'Вы не являетесь согласующим для этого документа'
            }, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def reject(self, request, pk=None):
        """Отклонить документ"""
        document = self.get_object()
        user = request.user

        try:
            approval = Approval.objects.get(
                document=document,
                approver=user,
                decision='pending'
            )
            if not _is_current_step(document, approval):
                return Response({
                    'status': 'error',
                    'message': 'Сейчас не ваш шаг согласования'
                }, status=status.HTTP_400_BAD_REQUEST)
            approval.decision = 'rejected'
            approval.decided_at = timezone.now()
            approval.comment = request.data.get('comment', '')
            approval.save()

            Approval.objects.filter(
                document=document,
                decision='pending'
            ).exclude(id=approval.id).update(
                decision='returned',
                decided_at=approval.decided_at
            )

            returned_status = (
                DocumentStatus.objects.filter(name__iexact='Возвращено на доработку').first()
                or DocumentStatus.objects.filter(name__iexact='На доработке').first()
                or DocumentStatus.objects.filter(name__iexact='Черновик').first()
            )
            if returned_status:
                document.status = returned_status
                document.last_rejection_comment = approval.comment
                document.last_rejection_at = approval.decided_at
                document.save(update_fields=['status', 'last_rejection_comment', 'last_rejection_at'])

            return Response({
                'status': 'success',
                'message': f'Документ {document.registration_number} отклонен'
            })
        except Approval.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'Вы не являетесь согласующим для этого документа'
            }, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def return_to_author(self, request, pk=None):
        """Вернуть документ на доработку после выполнения"""
        document = self.get_object()
        user = request.user

        approval = (
            Approval.objects
            .filter(document=document, approver=user)
            .order_by('-cycle', '-created_at')
            .first()
        )
        if not approval:
            return Response({
                'status': 'error',
                'message': 'Вы не являетесь согласующим для этого документа'
            }, status=status.HTTP_400_BAD_REQUEST)

        approval.decision = 'returned'
        approval.decided_at = timezone.now()
        approval.comment = request.data.get('comment', '')
        if not approval.comment.strip():
            return Response({
                'status': 'error',
                'message': 'Комментарий обязателен для возврата на доработку'
            }, status=status.HTTP_400_BAD_REQUEST)
        approval.save()

        Approval.objects.filter(
            document=document,
            decision='pending',
            cycle=approval.cycle
        ).exclude(id=approval.id).update(
            decision='returned',
            decided_at=approval.decided_at
        )

        returned_status = (
            DocumentStatus.objects.filter(name__iexact='Возвращено на доработку').first()
            or DocumentStatus.objects.filter(name__iexact='На доработке').first()
            or DocumentStatus.objects.filter(name__iexact='Черновик').first()
        )
        if returned_status:
            document.status = returned_status
            document.last_rejection_comment = approval.comment
            document.last_rejection_at = approval.decided_at
            document.actual_deadline = None
            document.save(update_fields=['status', 'last_rejection_comment', 'last_rejection_at', 'actual_deadline'])

        files = []
        if hasattr(request, 'FILES'):
            files = request.FILES.getlist('files')

        for file_obj in files:
            doc_file = DocumentFile.objects.create(
                document=document,
                file=file_obj,
                uploaded_by=user,
            )
            DocumentVersion.objects.create(
                document=document,
                file=doc_file.file,
                created_by=user,
            )

        if document.author:
            Notification.objects.create(
                user=document.author,
                notification_type='status_change',
                title='Документ возвращен на доработку',
                text=f'{document.registration_number} — {document.title}',
                link=f'/documents/outgoing/?open={document.id}',
                document=document,
                sender=user
            )

        return Response({
            'status': 'success',
            'message': f'Документ {document.registration_number} возвращен на доработку'
        })

    @action(detail=True, methods=['post'])
    def acknowledge(self, request, pk=None):
        """Ознакомиться с документом"""
        document = self.get_object()
        user = request.user

        try:
            approval = Approval.objects.get(
                document=document,
                approver=user,
                decision='pending'
            )
            if not _is_current_step(document, approval):
                return Response({
                    'status': 'error',
                    'message': 'Сейчас не ваш шаг согласования'
                }, status=status.HTTP_400_BAD_REQUEST)
            approval.decision = 'acknowledged'
            approval.decided_at = timezone.now()
            approval.comment = request.data.get('comment', '')
            approval.save()
            _maybe_set_actual_deadline(document)

            return Response({
                'status': 'success',
                'message': f'Документ {document.registration_number} отмечен как ознакомленный'
            })
        except Approval.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'Вы не являетесь согласующим для этого документа'
            }, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def execute(self, request, pk=None):
        """Отметить документ как исполненный"""
        document = self.get_object()
        user = request.user

        try:
            approval = Approval.objects.get(
                document=document,
                approver=user,
                decision='pending'
            )
            if not _is_current_step(document, approval):
                return Response({
                    'status': 'error',
                    'message': 'Сейчас не ваш шаг согласования'
                }, status=status.HTTP_400_BAD_REQUEST)
            approval.decision = 'executed'
            approval.decided_at = timezone.now()
            approval.comment = request.data.get('comment', '')
            approval.save()
            _maybe_set_actual_deadline(document)

            return Response({
                'status': 'success',
                'message': f'Документ {document.registration_number} отмечен как исполненный'
            })
        except Approval.DoesNotExist:
            return Response({
                'status': 'error',
                'message': 'Вы не являетесь согласующим для этого документа'
            }, status=status.HTTP_400_BAD_REQUEST)

    @action(detail=True, methods=['post'])
    def resubmit(self, request, pk=None):
        """Повторно отправить документ после отклонения"""
        document = self.get_object()
        user = request.user

        if not user.is_authenticated or document.author != user:
            return Response({
                'status': 'error',
                'message': 'Только автор может повторно отправить документ'
            }, status=status.HTTP_403_FORBIDDEN)

        action_type = request.data.get('action_type') or document.action_type
        action_status_map = {
            'approve': 'На согласовании',
            'acknowledge': 'На ознакомлении',
            'execute': 'На исполнении',
        }
        status_name = action_status_map.get(action_type, 'На согласовании')
        doc_status = DocumentStatus.objects.filter(name__iexact=status_name).first()
        if not doc_status:
            doc_status, _ = DocumentStatus.objects.get_or_create(
                name=status_name,
                defaults={'color': '#6c757d', 'is_final': False}
            )
        if doc_status and getattr(doc_status, 'is_final', False):
            doc_status.is_final = False
            doc_status.save(update_fields=['is_final'])
        if not doc_status:
            doc_status = DocumentStatus.objects.order_by('id').first()
        if not doc_status:
            return Response({
                'status': 'error',
                'message': 'Не найден статус для повторной отправки'
            }, status=status.HTTP_400_BAD_REQUEST)

        if action_type in ['approve', 'acknowledge', 'execute']:
            document.action_type = action_type

        document.status = doc_status

        document.actual_deadline = None
        document.save(update_fields=['status', 'action_type', 'actual_deadline'])

        files = []
        if hasattr(request, 'FILES'):
            files = request.FILES.getlist('files')

        for file_obj in files:
            doc_file = DocumentFile.objects.create(
                document=document,
                file=file_obj,
                uploaded_by=user,
            )
            DocumentVersion.objects.create(
                document=document,
                file=doc_file.file,
                created_by=user,
            )

        max_cycle = Approval.objects.filter(document=document).aggregate(max_cycle=Max('cycle'))['max_cycle'] or 1
        next_cycle = max_cycle + 1

        delivery_mode = document.delivery_mode
        approval_order = request.data.get('approval_order') or document.approval_order or 'sequential'

        if delivery_mode == 'auto':
            template = (
                DocumentRouteTemplate.objects
                .filter(document_type=document.document_type, is_active=True)
                .annotate(step_count=Count('steps'))
                .filter(step_count__gt=0)
                .first()
            )

            if not template:
                return Response({
                    'status': 'error',
                    'message': 'Не найден шаблон маршрута для повторной отправки'
                }, status=status.HTTP_400_BAD_REQUEST)

            if not request.data.get('approval_order'):
                approval_order = template.approval_order
            if document.approval_order != approval_order:
                document.approval_order = approval_order
                document.save(update_fields=['approval_order'])
            sequential_created = []
            for step in template.steps.all().order_by('step_number'):
                users = []
                if step.user:
                    resolved = _resolve_approver(step.user)
                    if resolved and resolved != document.author:
                        users.append(resolved)
                elif step.department:
                    users.extend(
                        User.objects.filter(
                            department=step.department,
                            is_active=True
                        )
                    )
                users = [_resolve_approver(u) for u in users]
                users = [u for u in users if u]
                users = [u for u in users if u != document.author]
                users = list({u.id: u for u in users}.values())
                for approver in users:
                    approval, created = Approval.objects.get_or_create(
                        document=document,
                        approver=approver,
                        step=step.step_number if approval_order == 'sequential' else 1,
                        cycle=next_cycle,
                        defaults={
                            'decision': 'pending',
                            'is_required': True
                        }
                    )
                    if created:
                        if approval_order == 'sequential':
                            sequential_created.append(approval)
                        else:
                            Notification.objects.create(
                                user=approver,
                                notification_type='new_document',
                                title='Новый документ',
                                text=f'{document.registration_number} — {document.title}',
                                link=f'/documents/incoming/?open={document.id}',
                                document=document,
                                sender=user
                            )
            if approval_order == 'sequential' and sequential_created:
                min_step = min(a.step for a in sequential_created)
                for approval in [a for a in sequential_created if a.step == min_step]:
                    Notification.objects.create(
                        user=approval.approver,
                        notification_type='new_document',
                        title='Новый документ',
                        text=f'{document.registration_number} — {document.title}',
                        link=f'/documents/incoming/?open={document.id}',
                        document=document,
                        sender=user
                    )
        else:
            manual_route = request.data.get('manual_route', document.manual_route or [])
            if isinstance(manual_route, str):
                try:
                    manual_route = json.loads(manual_route)
                except json.JSONDecodeError:
                    manual_route = []

            if not manual_route:
                return Response({
                    'status': 'error',
                    'message': 'Для ручного маршрута нужен manual_route'
                }, status=status.HTTP_400_BAD_REQUEST)

            sequential_created = []
            for i, step_data in enumerate(manual_route, start=1):
                step_number = i if approval_order == 'sequential' else 1

                if step_data.get('type') == 'user':
                    try:
                        approver = User.objects.get(
                            id=step_data['id'],
                            is_active=True
                        )
                        approver = _resolve_approver(approver)
                        if not approver:
                            continue
                        approval, created = Approval.objects.get_or_create(
                            document=document,
                            approver=approver,
                            step=step_number,
                            cycle=next_cycle,
                            defaults={
                                'decision': 'pending',
                                'is_required': True
                            }
                        )
                        if created:
                            if approval_order == 'sequential':
                                sequential_created.append(approval)
                            else:
                                Notification.objects.create(
                                    user=approver,
                                    notification_type='new_document',
                                    title='Новый документ',
                                    text=f'{document.registration_number} — {document.title}',
                                    link=f'/documents/incoming/?open={document.id}',
                                    document=document,
                                    sender=user
                                )
                    except User.DoesNotExist:
                        continue
                elif step_data.get('type') == 'department':
                    users = User.objects.filter(
                        department_id=step_data['id'],
                        is_active=True
                    )
                    resolved_users = [_resolve_approver(u) for u in users]
                    resolved_users = [u for u in resolved_users if u]
                    resolved_users = list({u.id: u for u in resolved_users}.values())
                    for approver in resolved_users:
                        approval, created = Approval.objects.get_or_create(
                            document=document,
                            approver=approver,
                            step=step_number,
                            cycle=next_cycle,
                            defaults={
                                'decision': 'pending',
                                'is_required': True
                            }
                        )
                        if created:
                            if approval_order == 'sequential':
                                sequential_created.append(approval)
                            else:
                                Notification.objects.create(
                                    user=approver,
                                    notification_type='new_document',
                                    title='Новый документ',
                                    text=f'{document.registration_number} — {document.title}',
                                    link=f'/documents/incoming/?open={document.id}',
                                    document=document,
                                    sender=user
                                )
            if approval_order == 'sequential' and sequential_created:
                min_step = min(a.step for a in sequential_created)
                for approval in [a for a in sequential_created if a.step == min_step]:
                    Notification.objects.create(
                        user=approval.approver,
                        notification_type='new_document',
                        title='Новый документ',
                        text=f'{document.registration_number} — {document.title}',
                        link=f'/documents/incoming/?open={document.id}',
                        document=document,
                        sender=user
                    )

        return Response({
            'status': 'success',
            'message': 'Документ повторно отправлен'
        })

    @action(detail=True, methods=['post'])
    def archive(self, request, pk=None):
        """Архивировать документ (только автор и только финальный статус)"""
        document = self.get_object()
        user = request.user

        if not user.is_authenticated or document.author != user:
            return Response({
                'status': 'error',
                'message': 'Только автор может архивировать документ'
            }, status=status.HTTP_403_FORBIDDEN)

        if document.is_archived:
            return Response({
                'status': 'success',
                'message': 'Документ уже в архиве'
            })

        if not document.status or not document.status.is_final:
            return Response({
                'status': 'error',
                'message': 'Документ можно архивировать только после завершения'
            }, status=status.HTTP_400_BAD_REQUEST)

        archive_status = (
            DocumentStatus.objects.filter(name__iexact='Архив').first()
            or DocumentStatus.objects.filter(name__iexact='В архиве').first()
        )
        if not archive_status:
            archive_status, _ = DocumentStatus.objects.get_or_create(
                name='Архив',
                defaults={'color': '#6c757d', 'is_final': True}
            )

        document.is_archived = True
        if archive_status:
            document.status = archive_status
            document.save(update_fields=['is_archived', 'status'])
        else:
            document.save(update_fields=['is_archived'])

        return Response({
            'status': 'success',
            'message': f'Документ {document.registration_number} перенесен в архив'
        })

    @action(detail=True, methods=['post'])
    def unarchive(self, request, pk=None):
        """Вернуть документ из архива (только автор)"""
        document = self.get_object()
        user = request.user

        if not user.is_authenticated or document.author != user:
            return Response({
                'status': 'error',
                'message': 'Только автор может вернуть документ из архива'
            }, status=status.HTTP_403_FORBIDDEN)

        if not document.is_archived:
            return Response({
                'status': 'success',
                'message': 'Документ не в архиве'
            })

        document.is_archived = False
        document.save(update_fields=['is_archived'])

        return Response({
            'status': 'success',
            'message': f'Документ {document.registration_number} возвращен из архива'
        })


def _maybe_set_actual_deadline(document):
    _ensure_cycle(document=document)
    max_cycle = Approval.objects.filter(document=document).aggregate(max_cycle=Max('cycle'))['max_cycle']
    if not max_cycle:
        return
    approvals = Approval.objects.filter(document=document, cycle=max_cycle)
    if not approvals.exists():
        return
    if approvals.filter(decision='pending').exists():
        return
    if approvals.filter(decision__in=['rejected', 'returned']).exists():
        return
    status_map = {
        'approve': ['Согласовано', 'Согласован'],
        'acknowledge': ['Ознакомлено', 'Ознакомлен'],
        'execute': ['Исполнено', 'Исполнен'],
    }
    desired_names = status_map.get(document.action_type, [])
    final_status = None
    for name in desired_names:
        final_status = DocumentStatus.objects.filter(name__iexact=name).first()
        if final_status:
            break
        final_status, _ = DocumentStatus.objects.get_or_create(
            name=name,
            defaults={'color': '#6c757d', 'is_final': True}
        )
        if final_status:
            break
    if not final_status:
        final_status = DocumentStatus.objects.filter(is_final=True).order_by('id').first()
    if final_status and document.status_id != final_status.id:
        document.status = final_status
    if not document.actual_deadline:
        document.actual_deadline = timezone.now().date()
        document.save(update_fields=['actual_deadline', 'status'])
    else:
        document.save(update_fields=['status'])


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def DocumentHistoryAPIView(request, pk):
    document = Document.objects.filter(pk=pk, author=request.user).first()
    if not document:
        return Response({
            'status': 'error',
            'message': 'Документ не найден или нет доступа'
        }, status=status.HTTP_404_NOT_FOUND)

    _ensure_cycle(document=document)
    approvals = (
        Approval.objects
        .filter(document=document)
        .select_related('approver')
        .order_by('cycle', 'step', 'created_at')
    )
    files = DocumentFile.objects.filter(document=document).order_by('uploaded_at')
    versions = DocumentVersion.objects.filter(document=document).order_by('version')

    lines = [
        f"Документ: {document.registration_number}",
        f"Тема: {document.title}",
        f"Тип: {document.document_type.name}",
        f"Автор: {document.author.full_name}",
        f"Статус: {document.status.name}",
        f"Создан: {document.created_at.strftime('%d.%m.%Y %H:%M')}",
        f"Срок: {document.deadline.strftime('%d.%m.%Y') if document.deadline else '—'}",
        f"Фактический срок: {document.actual_deadline.strftime('%d.%m.%Y') if document.actual_deadline else '—'}",
        f"Тип обработки: {document.get_action_type_display()}",
        f"Приоритет: {document.get_priority_display()}",
        f"Входящий номер: {document.external_number or '—'}",
        f"Входящая дата: {document.external_date.strftime('%d.%m.%Y') if document.external_date else '—'}",
        f"Корреспондент: {document.correspondent or '—'}",
        f"Описание: {document.description or '—'}",
        "",
        "История согласования:",
    ]

    for approval in approvals:
        decided_at = approval.decided_at.strftime('%d.%m.%Y %H:%M') if approval.decided_at else '—'
        comment = approval.comment or '—'
        lines.append(
            f"- Раунд {approval.cycle}, шаг {approval.step}: {approval.approver.full_name} — "
            f"{approval.get_decision_display()} ({decided_at})"
        )
        lines.append(f"  Комментарий: {comment}")

    lines.append("")
    lines.append("Файлы документа:")
    if files.exists():
        for f in files:
            uploaded_at = f.uploaded_at.strftime('%d.%m.%Y %H:%M') if f.uploaded_at else '—'
            lines.append(
                f"- {f.file_name or f.file.name} | {f.file_type or '—'} | "
                f"{f.file_size or 0} байт | {uploaded_at} | "
                f"Загрузил: {f.uploaded_by.full_name if f.uploaded_by else '—'}"
            )
    else:
        lines.append("- —")

    lines.append("")
    lines.append("Версии документа:")
    if versions.exists():
        for v in versions:
            created_at = v.created_at.strftime('%d.%m.%Y %H:%M') if v.created_at else '—'
            lines.append(
                f"- Версия {v.version}: {v.file.name} | {created_at} | "
                f"Создал: {v.created_by.full_name if v.created_by else '—'}"
            )
    else:
        lines.append("- —")

    try:
        from docx import Document as DocxDocument
    except Exception:
        content = "\n".join(lines)
        response = HttpResponse(content, content_type='text/plain; charset=utf-8')
        response['Content-Disposition'] = f'attachment; filename="history_{document.registration_number}.txt"'
        return response

    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    docx = DocxDocument()
    normal_style = docx.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)

    title = docx.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("ОТЧЕТ О ДОКУМЕНТЕ")
    title_run.bold = True
    title_run.font.size = Pt(16)

    subtitle = docx.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(f"{document.registration_number} — {document.title}")
    subtitle_run.font.size = Pt(12)

    meta = docx.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Сформировано: {timezone.now().strftime('%d.%m.%Y %H:%M')}")

    docx.add_paragraph("")

    docx.add_heading("Основные сведения", level=2)
    info_table = docx.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    info_rows = [
        ("Регистрационный номер", document.registration_number),
        ("Тема документа", document.title),
        ("Тип документа", document.document_type.name),
        ("Статус", document.status.name),
        ("Тип обработки", document.get_action_type_display()),
        ("Приоритет", document.get_priority_display()),
        ("Автор", document.author.full_name),
        ("Создан", document.created_at.strftime('%d.%m.%Y %H:%M')),
        ("Срок", document.deadline.strftime('%d.%m.%Y') if document.deadline else '—'),
        ("Фактический срок", document.actual_deadline.strftime('%d.%m.%Y') if document.actual_deadline else '—'),
        ("Входящий номер", document.external_number or '—'),
        ("Входящая дата", document.external_date.strftime('%d.%m.%Y') if document.external_date else '—'),
        ("Корреспондент", document.correspondent or '—'),
        ("Описание", document.description or '—'),
    ]
    for label, value in info_rows:
        row = info_table.add_row().cells
        row[0].text = label
        row[1].text = value

    docx.add_paragraph("")
    docx.add_heading("История согласования", level=2)
    if approvals.exists():
        table = docx.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Раунд"
        hdr[1].text = "Шаг"
        hdr[2].text = "Согласующий"
        hdr[3].text = "Решение"
        hdr[4].text = "Дата"
        hdr[5].text = "Комментарий"
        for approval in approvals:
            decided_at = approval.decided_at.strftime('%d.%m.%Y %H:%M') if approval.decided_at else '—'
            comment = approval.comment or '—'
            row = table.add_row().cells
            row[0].text = str(approval.cycle)
            row[1].text = str(approval.step)
            row[2].text = approval.approver.full_name
            row[3].text = approval.get_decision_display()
            row[4].text = decided_at
            row[5].text = comment
    else:
        docx.add_paragraph("—")

    docx.add_heading("Файлы документа", level=2)
    if files.exists():
        table = docx.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Имя"
        hdr[1].text = "Тип"
        hdr[2].text = "Размер"
        hdr[3].text = "Дата загрузки"
        hdr[4].text = "Загрузил"
        for f in files:
            uploaded_at = f.uploaded_at.strftime('%d.%m.%Y %H:%M') if f.uploaded_at else '—'
            row = table.add_row().cells
            row[0].text = f.file_name or f.file.name
            row[1].text = f.file_type or '—'
            row[2].text = str(f.file_size or 0)
            row[3].text = uploaded_at
            row[4].text = f.uploaded_by.full_name if f.uploaded_by else '—'
    else:
        docx.add_paragraph("—")

    docx.add_heading("Версии документа", level=2)
    if versions.exists():
        table = docx.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = "Версия"
        hdr[1].text = "Файл"
        hdr[2].text = "Дата"
        hdr[3].text = "Создал"
        for v in versions:
            created_at = v.created_at.strftime('%d.%m.%Y %H:%M') if v.created_at else '—'
            row = table.add_row().cells
            row[0].text = str(v.version)
            row[1].text = v.file.name
            row[2].text = created_at
            row[3].text = v.created_by.full_name if v.created_by else '—'
    else:
        docx.add_paragraph("—")

    buffer = io.BytesIO()
    docx.save(buffer)
    buffer.seek(0)
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="history_{document.registration_number}.docx"'
    return response


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def DashboardStatsAPIView(request):
    user = request.user
    if not user.is_authenticated:
        return Response({
            'incoming_count': 0,
            'incoming_urgent': 0,
            'approval_count': 0,
            'overdue_count': 0,
            'approval_overdue': 0,
            'processed_today': 0,
        })

    today = timezone.localdate()

    min_pending_step = (
        Approval.objects.filter(
            document_id=OuterRef('document_id'),
            decision='pending',
        )
        .values('document_id')
        .annotate(min_step=Min('step'))
        .values('min_step')
    )

    pending_approvals = (
        Approval.objects
        .filter(approver=user, decision='pending', document__is_archived=False)
        .annotate(min_step=Subquery(min_pending_step))
        .filter(
            models.Q(document__action_type='acknowledge')
            | models.Q(document__approval_order='parallel')
            | models.Q(step=models.F('min_step'))
        )
    )

    incoming_doc_ids = pending_approvals.values('document_id').distinct()
    incoming_count = incoming_doc_ids.count()
    incoming_urgent = pending_approvals.filter(document__priority='urgent').values('document_id').distinct().count()
    overdue_count = pending_approvals.filter(document__deadline__lt=today).values('document_id').distinct().count()

    approval_count = (
        Document.objects
        .filter(author=user, is_archived=False)
        .exclude(status__is_final=True)
        .exclude(status__name__iexact='Черновик')
        .exclude(status__name__icontains='доработ')
        .exclude(status__name__icontains='возвращ')
        .count()
    )

    processed_today = Approval.objects.filter(
        approver=user,
        decided_at__date=today,
        decision__in=['approved', 'rejected', 'acknowledged', 'executed']
    ).count()

    return Response({
        'incoming_count': incoming_count,
        'incoming_urgent': incoming_urgent,
        'approval_count': approval_count,
        'overdue_count': overdue_count,
        'approval_overdue': overdue_count,
        'processed_today': processed_today,
    })


@api_view(['GET'])
@permission_classes([IsAuthenticated])
def NotificationsAPIView(request):
    notifications = (
        Notification.objects
        .filter(user=request.user)
        .order_by('-created_at')[:20]
    )
    data = [
        {
            'id': n.id,
            'title': n.title,
            'text': n.text,
            'link': n.link,
            'created_at': n.created_at.isoformat(),
            'is_read': n.is_read,
            'type': n.notification_type,
        }
        for n in notifications
    ]
    unread_count = Notification.objects.filter(user=request.user, is_read=False).count()
    return Response({
        'unread_count': unread_count,
        'results': data
    })


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def MarkAllNotificationsReadAPIView(request):
    Notification.objects.filter(user=request.user, is_read=False).update(
        is_read=True,
        read_at=timezone.now()
    )
    return Response({'status': 'success'})


@api_view(['POST'])
@permission_classes([IsAuthenticated])
def MarkNotificationReadAPIView(request, pk):
    notification = Notification.objects.filter(id=pk, user=request.user).first()
    if not notification:
        return Response({'status': 'error', 'message': 'Not found'}, status=status.HTTP_404_NOT_FOUND)
    if not notification.is_read:
        notification.is_read = True
        notification.read_at = timezone.now()
        notification.save(update_fields=['is_read', 'read_at'])
    return Response({'status': 'success'})


def _ensure_cycle(document=None):
    if document:
        Approval.objects.filter(document=document, cycle__isnull=True).update(cycle=1)
    else:
        Approval.objects.filter(cycle__isnull=True).update(cycle=1)


from rest_framework.viewsets import ReadOnlyModelViewSet
from .serializers import DocumentRouteSerializer


class DocumentRouteViewSet(ReadOnlyModelViewSet):
    serializer_class = DocumentRouteSerializer

    def get_queryset(self):
        queryset = DocumentRouteTemplate.objects.prefetch_related(
            'steps__user',
            'steps__department'
        )

        document_type = self.request.query_params.get('document_type')
        if document_type:
            queryset = queryset.filter(document_type_id=document_type)

        return queryset
